#!/usr/bin/env python3
"""
文本提取工具 - extract_text.py
Source: clinical-note-writer-v2.md (Scripts Section)

支持从 .txt / .pdf / .docx 文件中提取纯文本，
统一输出为字符串，供病历生成流程使用。
"""

import os
import sys
import logging
from pathlib import Path
from typing import Optional, Tuple

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s'
)
logger = logging.getLogger(__name__)


class TextExtractor:
    """多格式文本提取器"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}
    
    def __init__(self):
        self._pdf_backend = None  # 延迟加载
    
    def extract(self, file_path: str) -> Tuple[bool, str, str]:
        """
        从文件提取文本。
        
        Args:
            file_path: 文件路径（绝对或相对）
            
        Returns:
            (success: bool, text: str, message: str)
            success=True 表示成功提取，text为文本内容
        """
        path = Path(file_path)
        
        # 1. 文件存在性检查
        if not path.exists():
            return False, "", f"文件不存在: {file_path}"
        
        # 2. 扩展名检查
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, (
                ""
            ), f"不支持的文件格式: {ext}，支持的格式: {self.SUPPORTED_EXTENSIONS}"
        
        try:
            # 3. 按类型分派提取方法
            extractors = {
                '.txt': self._extract_txt,
                '.pdf': self._extract_pdf,
                '.docx': self._extract_docx,
                '.doc': self._extract_docx,  # 尝试用 docx 库处理（可能失败）
            }
            
            extractor = extractors.get(ext)
            if extractor is None:
                return False, "", f"无对应的提取器: {ext}"
            
            text = extractor(path)
            
            # 4. 后处理：清理空白
            text = self._postprocess(text)
            
            if len(text.strip()) < 10:
                logger.warning(f"文件内容过短（<10字符），可能是空文件或扫描件: {file_path}")
                return True, text, f"提取成功但内容较短（{len(text)}字符），请确认是否为扫描图片"
            
            logger.info(f"成功提取文本: {file_path} ({len(text)} 字符)")
            return True, text, "OK"
            
        except Exception as e:
            logger.error(f"提取失败 [{ext}]: {file_path} → {str(e)}")
            return False, "", f"提取失败: {str(e)}"
    
    def _extract_txt(self, path: Path) -> str:
        """提取纯文本文件"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                text = path.read_text(encoding=encoding)
                logger.debug(f"TXT 编码检测成功: {encoding}")
                return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # 所有编码都失败，用 utf-8 忽略错误
        return path.read_text(encoding='utf-8', errors='ignore')
    
    def _extract_pdf(self, path: Path) -> str:
        """提取 PDF 文本"""
        text = ""
        
        # 策略1: pdfplumber（表格支持更好）
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            if text.strip():
                return text
        except ImportError:
            logger.debug("pdfplumber 未安装，尝试 PyPDF2")
        except Exception as e:
            logger.warning(f"pdfplumber 提取异常: {e}")
        
        # 策略2: PyPDF2
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(str(path))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            if text.strip():
                return text
        except ImportError:
            logger.debug("PyPDF2 未安装")
        except Exception as e:
            logger.warning(f"PyPDF2 提取异常: {e}")
        
        raise ImportError(
            "PDF 提取需要安装 pdfplumber 或 PyPDF2。"
            "请运行: pip install pdfplumber"
        )
    
    def _extract_docx(self, path: Path) -> str:
        """提取 DOCX/DOC 文件"""
        try:
            from docx import Document
            doc = Document(str(path))
            
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            
            # 同时提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))
            
            text = "\n".join(paragraphs)
            return text
            
        except ImportError:
            raise ImportError(
                "DOCX 提取需要安装 python-docx。"
                "请运行: pip install python-docx"
            )
    
    @staticmethod
    def _postprocess(text: str) -> str:
        """后处理：清理多余空白"""
        import re
        
        # 合并连续空白行（保留最多2个换行）
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 清理行尾空白
        text = '\n'.join(line.rstrip() for line in text.split('\n'))
        # 去除首尾空白
        text = text.strip()
        
        return text


# ============================================================
# CLI 接口
# ============================================================

def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='从 txt/pdf/docx 文件提取纯文本',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  python extract_text.py patient_record.pdf
  python extract_text.py admission.docx -o output.txt
  python extract_text.py note.txt --quiet
        """
    )
    parser.add_argument('input_file', help='输入文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径（默认输出到stdout）')
    parser.add_argument('-q', '--quiet', action='store_true', help='静默模式')
    
    args = parser.parse_args()
    
    if args.quiet:
        logger.setLevel(logging.ERROR)
    
    extractor = TextExtractor()
    success, text, message = extractor.extract(args.input_file)
    
    if not success:
        print(f"❌ 错误: {message}", file=sys.stderr)
        sys.exit(1)
    
    if args.output:
        output_path = Path(args.output)
        output_path.write_text(text, encoding='utf-8')
        print(f"✅ 文本已写入: {output_path} ({len(text)} 字符)")
    else:
        print(text)


if __name__ == '__main__':
    main()
